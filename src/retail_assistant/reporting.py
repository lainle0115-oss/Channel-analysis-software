from __future__ import annotations

from io import BytesIO

import pandas as pd
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from .analytics import detect_anomalies, latest_inventory_rows


WEEKDAY_NAMES = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
FIXED_RETAIL_DATES = {
    (1, 1): "元旦",
    (2, 14): "情人节",
    (3, 8): "妇女节",
    (5, 1): "劳动节",
    (6, 1): "儿童节",
    (10, 1): "国庆节",
    (11, 11): "双十一",
    (12, 12): "双十二",
    (12, 25): "圣诞节",
}

PIPELINE_FIELDS = [
    "Account / Company",
    "Contact Person",
    "Lead Source",
    "Stage",
    "Deal Size",
    "Probability",
    "Expected Close Date",
    "Pain Point",
    "Decision Maker",
    "Next Step",
    "Risk",
    "Status Update",
]
PIPELINE_STAGE_ORDER = ["初步沟通", "技术交流", "测试", "报价", "谈判", "合同/采购", "成交", "流失"]

COLUMN_LABELS = {
    "date": "日期",
    "week_start": "周开始日期",
    "channel": "渠道",
    "sku_id": "SKU 编码",
    "sku_name": "商品名称",
    "sales_qty": "销量",
    "sales_amount": "销售额",
    "purchase_qty": "采购数量",
    "purchase_sales_ratio": "采销比",
    "purchase_sales_gap": "采购-销售差额",
    "stock_qty": "库存",
    "latest_sales_qty": "最新销量",
    "latest_stock": "最新库存",
    "active_days": "有效天数",
    "previous_sales_qty": "上期销量",
    "previous_sales_amount": "上期销售额",
    "yoy_sales_qty": "去年同期销量",
    "yoy_sales_amount": "去年同期销售额",
    "yoy_qty_change": "销量同比变化",
    "current_week_qty": "本周销量",
    "current_week_amount": "本周销售额",
    "current_week_days": "本周有效天数",
    "previous_week_qty": "上周销量",
    "previous_week_amount": "上周销售额",
    "previous_week_days": "上周有效天数",
    "forecast_week_qty": "本周预测销量",
    "week_qty_change": "周销量变化",
    "qty_change": "销量变化",
    "forecast_qty": "周期预测销量",
    "remaining_days": "剩余天数",
    "avg_sales_qty": "历史日均销量",
    "sales_change": "销量变化",
    "issue": "问题",
    "action": "建议动作",
    "reason": "原因判断",
    "analysis": "分析",
}


def _chinese_table(frame: pd.DataFrame) -> pd.DataFrame:
    return frame.rename(columns=COLUMN_LABELS)


def _empty_frame(columns: list[str]) -> pd.DataFrame:
    return pd.DataFrame(columns=columns)


def _with_purchase_column(data: pd.DataFrame) -> pd.DataFrame:
    if "purchase_qty" in data:
        return data
    result = data.copy()
    result["purchase_qty"] = 0
    return result


def _context_value(context: dict[str, object] | None, key: str, default: str = "未提供") -> str:
    if not context:
        return default
    value = str(context.get(key, "")).strip()
    return value or default


def _action_register(anomalies: pd.DataFrame, report_end: pd.Timestamp) -> pd.DataFrame:
    columns = ["优先级", "渠道", "SKU", "问题", "事实依据", "建议动作", "负责人", "截止日期", "处理状态"]
    if anomalies.empty:
        return pd.DataFrame(columns=columns)
    rows = []
    for row in anomalies.head(20).itertuples(index=False):
        priority = "P0" if row.issue == "缺货风险" else "P1" if row.issue == "库存覆盖不足" else "P2"
        due_days = 1 if priority == "P0" else 2 if priority == "P1" else 3
        rows.append({
            "优先级": priority,
            "渠道": row.channel,
            "SKU": row.sku_name,
            "问题": row.issue,
            "事实依据": f"销量 {row.sales_qty:,.0f} 件｜库存 {row.stock_qty:,.0f} 件｜较历史均值 {row.sales_change:+.1%}",
            "建议动作": row.action,
            "负责人": "待分配",
            "截止日期": f"{report_end + pd.Timedelta(days=due_days):%Y-%m-%d}",
            "处理状态": "待处理",
        })
    return pd.DataFrame(rows, columns=columns)


def _style_workbook(writer: pd.ExcelWriter) -> None:
    header_fill = PatternFill("solid", fgColor="17324D")
    header_font = Font(color="FFFFFF", bold=True)
    accent_fill = PatternFill("solid", fgColor="E8F1F8")
    for sheet in writer.book.worksheets:
        if sheet.title == "领导摘要":
            continue
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        sheet.sheet_view.showGridLines = False
        sheet.page_setup.orientation = "landscape"
        sheet.page_setup.fitToWidth = 1
        sheet.sheet_properties.pageSetUpPr.fitToPage = True
        for cell in sheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(vertical="center")
        sheet.row_dimensions[1].height = 24
        for row_index in range(2, min(sheet.max_row, 6) + 1):
            if row_index % 2 == 0:
                for cell in sheet[row_index]:
                    cell.fill = accent_fill
        for column_cells in sheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 2, 12), 56)
            for cell in column_cells:
                cell.alignment = Alignment(vertical="top", wrap_text=True)


def _write_leadership_summary(
    writer: pd.ExcelWriter,
    title: str,
    period_text: str,
    data: pd.DataFrame,
    report_text: str,
    actions: pd.DataFrame,
) -> None:
    workbook = writer.book
    sheet = workbook.create_sheet("领导摘要", 0)
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A16"
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 1
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.print_area = "A1:H45"

    navy = PatternFill("solid", fgColor="17324D")
    teal = PatternFill("solid", fgColor="168B83")
    pale = PatternFill("solid", fgColor="E8F1F8")
    white_font = Font(color="FFFFFF", bold=True)
    border = Border(bottom=Side(style="thin", color="DDE4E8"))

    sheet.merge_cells("A1:H2")
    sheet["A1"] = title
    sheet["A1"].font = Font(color="FFFFFF", bold=True, size=20)
    sheet["A1"].fill = navy
    sheet["A1"].alignment = Alignment(vertical="center")
    sheet.row_dimensions[1].height = 28
    sheet.row_dimensions[2].height = 28
    sheet.merge_cells("A3:H3")
    sheet["A3"] = f"报告周期：{period_text}｜生成后可直接用于经营复盘和行动跟进"
    sheet["A3"].font = Font(color="647481", italic=True)

    latest = latest_inventory_rows(data)
    cards = [
        ("本期销量", f"{data['sales_qty'].sum():,.0f} 件" if not data.empty else "0 件"),
        ("本期销售额", f"¥{data['sales_amount'].sum():,.0f}" if not data.empty and data["sales_amount"].abs().sum() else "未提供"),
        ("渠道 / 商品", f"{data['channel'].nunique()} / {data['sku_id'].nunique()}" if not data.empty else "0 / 0"),
        ("最新库存 / 待办", f"{latest['stock_qty'].sum():,.0f} / {len(actions)}" if not data.empty else f"0 / {len(actions)}"),
    ]
    for index, (label, value) in enumerate(cards):
        start_col = index * 2 + 1
        sheet.merge_cells(start_row=5, start_column=start_col, end_row=5, end_column=start_col + 1)
        sheet.merge_cells(start_row=6, start_column=start_col, end_row=7, end_column=start_col + 1)
        label_cell = sheet.cell(5, start_col, label)
        value_cell = sheet.cell(6, start_col, value)
        label_cell.fill = pale
        label_cell.font = Font(color="647481", bold=True)
        value_cell.fill = pale
        value_cell.font = Font(color="17324D", bold=True, size=18)
        value_cell.alignment = Alignment(vertical="center")

    sheet.merge_cells("A9:H9")
    sheet["A9"] = "管理层结论"
    sheet["A9"].fill = teal
    sheet["A9"].font = white_font
    conclusion_lines = [
        line[2:].replace("**", "")
        for line in report_text.splitlines()
        if line.startswith("- ")
    ][:4]
    for row_index, line in enumerate(conclusion_lines, start=10):
        sheet.merge_cells(start_row=row_index, start_column=1, end_row=row_index, end_column=8)
        sheet.cell(row_index, 1, f"• {line}")
        sheet.cell(row_index, 1).alignment = Alignment(vertical="top", wrap_text=True)
        sheet.cell(row_index, 1).border = border
        sheet.row_dimensions[row_index].height = 34

    sheet.merge_cells("A15:H15")
    sheet["A15"] = "重点行动跟进"
    sheet["A15"].fill = teal
    sheet["A15"].font = white_font
    action_headers = ["优先级", "渠道", "商品", "问题", "建议动作", "负责人", "截止日期", "处理状态"]
    for column, header in enumerate(action_headers, start=1):
        cell = sheet.cell(16, column, header)
        cell.fill = navy
        cell.font = white_font
        cell.alignment = Alignment(vertical="center")
    for row_index, row in enumerate(actions.head(10).itertuples(index=False), start=17):
        values = [row.优先级, row.渠道, row.SKU, row.问题, row.建议动作, row.负责人, row.截止日期, row.处理状态]
        for column, value in enumerate(values, start=1):
            cell = sheet.cell(row_index, column, value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
        sheet.row_dimensions[row_index].height = 36
    if actions.empty:
        sheet.merge_cells("A17:H17")
        sheet["A17"] = "当前没有需要跟进的异常事项。"
        sheet["A17"].alignment = Alignment(vertical="center")
        sheet["A17"].fill = pale
        sheet.row_dimensions[17].height = 36

    widths = [12, 16, 24, 18, 42, 14, 14, 14]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[chr(64 + index)].width = width

    if "渠道对比" in workbook.sheetnames:
        source = workbook["渠道对比"]
        if source.max_row >= 2 and source.max_column >= 2:
            chart = BarChart()
            chart.type = "bar"
            chart.style = 10
            chart.title = "渠道销量对比"
            chart.x_axis.title = "销量"
            chart.x_axis.scaling.min = 0
            chart.legend = None
            chart.height = 7.5
            chart.width = 20
            chart.add_data(Reference(source, min_col=2, min_row=1, max_row=min(source.max_row, 9)), titles_from_data=True)
            chart.set_categories(Reference(source, min_col=1, min_row=2, max_row=min(source.max_row, 9)))
            for row_index in range(27, 46):
                sheet.row_dimensions[row_index].height = 20
            sheet.add_chart(chart, "A27")


def calendar_context(report_date: pd.Timestamp, context: dict[str, object] | None = None) -> dict[str, str]:
    report_date = pd.Timestamp(report_date)
    holiday = _context_value(context, "holiday_note", "")
    if not holiday:
        holiday = FIXED_RETAIL_DATES.get((report_date.month, report_date.day), "非内置节假日")
    date_tags = [WEEKDAY_NAMES[report_date.weekday()]]
    if report_date.weekday() >= 5:
        date_tags.append("周末")
    if report_date.day <= 3:
        date_tags.append("月初")
    if report_date.day >= 25:
        date_tags.append("月末")
    return {
        "日期": f"{report_date:%Y-%m-%d}",
        "星期": WEEKDAY_NAMES[report_date.weekday()],
        "日期标签": "、".join(date_tags),
        "节假日": holiday,
        "天气": _context_value(context, "weather_note"),
        "渠道活动": _context_value(context, "activity_note"),
    }


def _pipeline_frame(context: dict[str, object] | None) -> pd.DataFrame:
    records = (context or {}).get("pipeline_records", [])
    if not isinstance(records, list) or not records:
        return _empty_frame([*PIPELINE_FIELDS, "weighted_amount"])
    frame = pd.DataFrame(records)
    if "Contact" in frame.columns and "Contact Person" not in frame.columns:
        frame["Contact Person"] = frame["Contact"]
    for column in PIPELINE_FIELDS:
        if column not in frame.columns:
            frame[column] = ""
    frame = frame[PIPELINE_FIELDS].copy()
    frame["Deal Size"] = pd.to_numeric(frame["Deal Size"], errors="coerce").fillna(0.0)
    probability = pd.to_numeric(frame["Probability"], errors="coerce").fillna(0.0)
    if probability.max() > 1:
        probability = probability / 100
    frame["Probability"] = probability.clip(0, 1)
    frame["Expected Close Date"] = pd.to_datetime(frame["Expected Close Date"], errors="coerce")
    for column in ("Account / Company", "Contact Person", "Lead Source", "Stage", "Pain Point", "Decision Maker", "Next Step", "Risk", "Status Update"):
        frame[column] = frame[column].fillna("").astype(str).str.strip()
    frame["weighted_amount"] = frame["Deal Size"] * frame["Probability"]
    return frame


def _pipeline_summary(context: dict[str, object] | None, report_end: pd.Timestamp) -> dict[str, object]:
    frame = _pipeline_frame(context)
    if frame.empty:
        return {
            "frame": frame,
            "open_frame": frame,
            "open_count": 0,
            "open_amount": 0.0,
            "weighted_amount": 0.0,
            "upcoming_count": 0,
            "risk_count": 0,
            "top_stage": "未维护",
        }
    open_frame = frame[~frame["Stage"].isin(["成交", "流失"])].copy()
    close_dates = pd.to_datetime(open_frame["Expected Close Date"], errors="coerce")
    upcoming = open_frame[
        close_dates.between(pd.Timestamp(report_end).normalize(), pd.Timestamp(report_end).normalize() + pd.Timedelta(days=30))
    ]
    stage_totals = (
        open_frame.groupby("Stage", as_index=False)["weighted_amount"]
        .sum()
        .sort_values("weighted_amount", ascending=False)
    )
    return {
        "frame": frame,
        "open_frame": open_frame,
        "open_count": int(len(open_frame)),
        "open_amount": float(open_frame["Deal Size"].sum()),
        "weighted_amount": float(open_frame["weighted_amount"].sum()),
        "upcoming_count": int(len(upcoming)),
        "risk_count": int(open_frame["Risk"].astype(str).str.strip().ne("").sum()),
        "top_stage": str(stage_totals.iloc[0]["Stage"]) if not stage_totals.empty else "未维护",
    }


def _pipeline_markdown_lines(context: dict[str, object] | None, report_end: pd.Timestamp) -> list[str]:
    summary = _pipeline_summary(context, report_end)
    open_frame = summary["open_frame"]
    if open_frame.empty:
        return ["- 当前未维护销售 Pipeline，建议补充客户公司、阶段、预计金额、成交概率和下一步动作。"]
    lines = [
        f"- 开放商机 {summary['open_count']} 个，预计订单金额 ¥{summary['open_amount']:,.0f}，加权预测金额 ¥{summary['weighted_amount']:,.0f}；金额最集中的阶段为 {summary['top_stage']}。",
        f"- 未来 30 天内预计成交 {summary['upcoming_count']} 个，高风险项 {summary['risk_count']} 条；建议优先推进高概率、高金额且有明确下一步动作的客户。",
    ]
    priority = (
        open_frame.sort_values(["weighted_amount", "Deal Size"], ascending=False)
        .head(5)
    )
    for _, row in priority.iterrows():
        close_date = row.get("Expected Close Date")
        close_text = f"{pd.Timestamp(close_date):%Y-%m-%d}" if pd.notna(close_date) else "未定"
        lines.append(
            f"- {row.get('Account / Company')}｜{row.get('Stage')}：预计金额 ¥{float(row.get('Deal Size', 0)):,.0f}，"
            f"成交概率 {float(row.get('Probability', 0)):.0%}，预计成交 {close_text}；下一步：{row.get('Next Step') or '待补充'}；风险：{row.get('Risk') or '无'}。"
        )
    return lines


def _daily_channel(data: pd.DataFrame) -> pd.DataFrame:
    if data.empty:
        return _empty_frame(["date", "channel", "sales_qty", "sales_amount", "purchase_qty", "stock_qty"])
    data = data.copy()
    data["date"] = pd.to_datetime(data["date"]).dt.normalize()
    return (
        data.groupby(["date", "channel"], as_index=False)
        .agg(
            sales_qty=("sales_qty", "sum"),
            sales_amount=("sales_amount", "sum"),
            purchase_qty=("purchase_qty", "sum"),
            stock_qty=("stock_qty", "sum"),
        )
        .sort_values(["date", "channel"])
    )


def _period_totals(data: pd.DataFrame, start: pd.Timestamp, end: pd.Timestamp) -> pd.DataFrame:
    period = data[(data["date"] >= start) & (data["date"] <= end)]
    if period.empty:
        return _empty_frame(["channel", "sales_qty", "sales_amount", "purchase_qty", "active_days"])
    period = period.copy()
    period["_report_day"] = pd.to_datetime(period["date"]).dt.normalize()
    return (
        period.groupby("channel", as_index=False)
        .agg(
            sales_qty=("sales_qty", "sum"),
            sales_amount=("sales_amount", "sum"),
            purchase_qty=("purchase_qty", "sum"),
            active_days=("_report_day", "nunique"),
        )
    )


def daily_report_tables(data: pd.DataFrame, report_date: pd.Timestamp) -> dict[str, pd.DataFrame]:
    data = _with_purchase_column(data)
    report_date = pd.Timestamp(report_date).normalize()
    data = data.copy()
    data["_report_day"] = pd.to_datetime(data["date"]).dt.normalize()
    current = data[data["_report_day"] == report_date]
    previous_dates = data.loc[data["_report_day"] < report_date, "_report_day"]
    previous_date = previous_dates.max() if not previous_dates.empty else None
    previous = data[data["_report_day"] == previous_date] if previous_date is not None else data.iloc[0:0]

    channel = (
        current.groupby("channel", as_index=False)
        .agg(
            sales_qty=("sales_qty", "sum"),
            sales_amount=("sales_amount", "sum"),
            purchase_qty=("purchase_qty", "sum"),
            stock_qty=("stock_qty", "sum"),
        )
        if not current.empty
        else _empty_frame(["channel", "sales_qty", "sales_amount", "purchase_qty", "stock_qty"])
    )
    previous_channel = (
        previous.groupby("channel", as_index=False)["sales_qty"]
        .sum()
        .rename(columns={"sales_qty": "previous_sales_qty"})
        if not previous.empty
        else _empty_frame(["channel", "previous_sales_qty"])
    )
    channel = channel.merge(previous_channel, on="channel", how="left").fillna({"previous_sales_qty": 0})
    channel["qty_change"] = (
        (channel["sales_qty"] - channel["previous_sales_qty"])
        / channel["previous_sales_qty"].replace(0, pd.NA)
    ).fillna(0)

    sku = (
        current.groupby(["channel", "sku_id", "sku_name"], as_index=False)
        .agg(
            sales_qty=("sales_qty", "sum"),
            sales_amount=("sales_amount", "sum"),
            purchase_qty=("purchase_qty", "sum"),
            stock_qty=("stock_qty", "sum"),
        )
        .sort_values("sales_qty", ascending=False)
        if not current.empty
        else _empty_frame(["channel", "sku_id", "sku_name", "sales_qty", "sales_amount", "purchase_qty", "stock_qty"])
    )

    week_start = report_date - pd.Timedelta(days=report_date.weekday())
    current_week = _period_totals(data, week_start, report_date).rename(
        columns={
            "sales_qty": "current_week_qty",
            "sales_amount": "current_week_amount",
            "purchase_qty": "current_week_purchase_qty",
            "active_days": "current_week_days",
        }
    )
    previous_week = _period_totals(data, week_start - pd.Timedelta(days=7), week_start - pd.Timedelta(days=1)).rename(
        columns={
            "sales_qty": "previous_week_qty",
            "sales_amount": "previous_week_amount",
            "purchase_qty": "previous_week_purchase_qty",
            "active_days": "previous_week_days",
        }
    )
    weekly = current_week.merge(previous_week, on="channel", how="outer")
    if weekly.empty:
        weekly = _empty_frame([
            "channel", "current_week_qty", "current_week_amount", "current_week_days",
            "current_week_purchase_qty", "previous_week_qty", "previous_week_amount", "previous_week_purchase_qty",
            "previous_week_days",
        ])
    numeric_weekly_columns = [
        "current_week_qty", "current_week_amount", "current_week_purchase_qty", "current_week_days",
        "previous_week_qty", "previous_week_amount", "previous_week_purchase_qty", "previous_week_days",
    ]
    for column in numeric_weekly_columns:
        if column in weekly:
            weekly[column] = pd.to_numeric(weekly[column], errors="coerce").fillna(0)
    elapsed_days = max((report_date - week_start).days + 1, 1)
    weekly["forecast_week_qty"] = weekly["current_week_qty"] / elapsed_days * 7
    weekly["week_qty_change"] = 0.0
    has_previous = weekly["previous_week_qty"] != 0
    weekly.loc[has_previous, "week_qty_change"] = (
        (weekly.loc[has_previous, "current_week_qty"] - weekly.loc[has_previous, "previous_week_qty"])
        / weekly.loc[has_previous, "previous_week_qty"]
    )

    daily_trend = _daily_channel(data[data["date"] <= report_date])
    weekly_history_source = data[data["date"] <= report_date].copy()
    if weekly_history_source.empty:
        weekly_history = _empty_frame(["week_start", "channel", "sales_qty", "sales_amount", "active_days"])
    else:
        weekly_history_source["week_start"] = (
            weekly_history_source["date"] - pd.to_timedelta(weekly_history_source["date"].dt.weekday, unit="D")
        )
        latest_weeks = sorted(weekly_history_source["week_start"].drop_duplicates())[-4:]
        weekly_history = (
            weekly_history_source[weekly_history_source["week_start"].isin(latest_weeks)]
            .groupby(["week_start", "channel"], as_index=False)
            .agg(
                sales_qty=("sales_qty", "sum"),
                sales_amount=("sales_amount", "sum"),
                purchase_qty=("purchase_qty", "sum"),
                active_days=("date", "nunique"),
            )
            .sort_values(["week_start", "channel"])
        )
    anomalies = detect_anomalies(data[data["date"] <= report_date])
    return {
        "channel": channel.sort_values("sales_qty", ascending=False),
        "sku": sku,
        "weekly": weekly.sort_values("forecast_week_qty", ascending=False) if not weekly.empty else weekly,
        "weekly_history": weekly_history,
        "daily_trend": daily_trend,
        "anomalies": anomalies,
    }


def _reason_for_channel(
    row: pd.Series,
    context: dict[str, str],
    comparison_label: str = "上一数据日",
) -> str:
    change = float(row.get("qty_change", 0) or 0)
    activity = context["渠道活动"]
    weather = context["天气"]
    holiday = context["节假日"]
    signals: list[str] = []
    previous_sales_qty = float(row.get("previous_sales_qty", 0) or 0)
    if previous_sales_qty == 0:
        signals.append(f"{comparison_label}无可比销量数据，暂不判断涨跌，优先补齐基线。")
    elif change >= 0.2:
        signals.append(f"较{comparison_label}明显增长，优先核对渠道活动、陈列资源和补货承接是否同步放大。")
    elif change <= -0.2:
        signals.append(f"较{comparison_label}明显下滑，优先检查断货、价格、上架状态、竞品活动和履约覆盖。")
    else:
        signals.append(f"较{comparison_label}波动较温和，重点看 SKU 结构变化和库存覆盖。")
    if activity != "未提供":
        signals.append(f"已纳入活动信息：{activity}。")
    if weather != "未提供":
        signals.append(f"天气因素：{weather}，适合结合品类季节性判断需求变化。")
    if holiday != "非内置节假日":
        signals.append(f"日期因素：{holiday}，建议对比同类节日前后销售节奏。")
    return "".join(signals)


def build_daily_report(
    data: pd.DataFrame,
    report_date: pd.Timestamp,
    context: dict[str, object] | None = None,
) -> dict[str, object]:
    report_date = pd.Timestamp(report_date)
    tables = daily_report_tables(data, report_date)
    current = data[data["date"] == report_date]
    context_rows = calendar_context(report_date, context)
    if current.empty:
        return {
            "title": f"渠道销售日报 | {report_date:%Y-%m-%d}",
            "context": context_rows,
            "summary": ["当日无数据，无法生成经营判断。"],
            "tables": tables,
            "analysis": _empty_frame(["channel", "analysis"]),
            "actions": ["补充当日数据后重新生成日报。"],
        }

    qty = current["sales_qty"].sum()
    amount = current["sales_amount"].sum()
    stock = current["stock_qty"].sum()
    channels = tables["channel"].copy()
    weekly = tables["weekly"].copy()
    anomalies = tables["anomalies"].copy()
    if not channels.empty:
        reasons = channels.apply(lambda row: _reason_for_channel(row, context_rows), axis=1)
        channels["reason"] = reasons
    analysis = channels[["channel", "sales_qty", "previous_sales_qty", "qty_change", "stock_qty", "reason"]].copy()

    top_channel = channels.iloc[0]["channel"] if not channels.empty else "无"
    forecast_qty = weekly["forecast_week_qty"].sum() if not weekly.empty else 0
    previous_week_qty = weekly["previous_week_qty"].sum() if not weekly.empty else 0
    week_change = (forecast_qty - previous_week_qty) / previous_week_qty if previous_week_qty else 0
    amount_channels = current.groupby("channel")["sales_amount"].apply(lambda values: values.abs().sum() > 0)
    if amount <= 0:
        amount_text = "销售额字段不完整，暂不作为完整金额口径"
    elif amount_channels.all():
        amount_text = f"销售额 ¥{amount:,.0f}"
    else:
        amount_text = f"已提供渠道销售额 ¥{amount:,.0f}，未覆盖全部渠道"
    summary = [
        f"当日覆盖 {current['channel'].nunique()} 个渠道、{current['sku_id'].nunique()} 个 SKU，销量 {qty:,.0f} 件，{amount_text}。",
        f"{top_channel} 为当日销量最高渠道；本周按当前节奏预测销量 {forecast_qty:,.0f} 件，较上一完整周 {week_change:+.1%}。",
        f"日期标签：{context_rows['日期标签']}；节假日：{context_rows['节假日']}；天气：{context_rows['天气']}；渠道活动：{context_rows['渠道活动']}。",
        f"库存口径显示最新库存 {stock:,.0f} 件，当前识别 {len(anomalies)} 条异常或补货提醒。",
    ]
    actions = [
        "把涨跌幅超过 20% 的渠道逐一核对活动、价格、上架和履约状态。",
        "对库存覆盖不足 SKU 先锁定补货量，再决定是否继续放量。",
        "若天气或渠道活动未填写，补充后重新生成日报，避免原因判断停留在销售数据表面。",
    ]
    if not anomalies.empty:
        first = anomalies.iloc[0]
        actions.insert(0, f"优先处理 {first['channel']}｜{first['sku_name']}：{first['action']}。")
    return {
        "title": f"渠道销售日报 | {report_date:%Y-%m-%d}",
        "context": context_rows,
        "summary": summary,
        "tables": tables,
        "analysis": analysis,
        "actions": actions,
    }


def generate_daily_report(
    data: pd.DataFrame,
    report_date: pd.Timestamp,
    context: dict[str, object] | None = None,
) -> str:
    report = build_daily_report(data, report_date, context)
    tables = report["tables"]
    lines = [f"# {report['title']}", "", "## 经营结论"]
    lines.extend(f"- {item}" for item in report["summary"])
    lines.extend(["", "## 日期、天气与活动背景"])
    for key, value in report["context"].items():
        lines.append(f"- {key}：{value}")
    lines.extend(["", "## 渠道涨跌原因"])
    analysis = report["analysis"]
    if analysis.empty:
        lines.append("- 暂无渠道分析。")
    else:
        for row in analysis.itertuples(index=False):
            lines.append(f"- {row.channel}：销量 {row.sales_qty:,.0f} 件，日环比 {row.qty_change:+.1%}。{row.reason}")
    lines.extend(["", "## 本周销量预测"])
    weekly = tables["weekly"]
    if weekly.empty:
        lines.append("- 暂无本周预测。")
    else:
        for row in weekly.itertuples(index=False):
            lines.append(
                f"- {row.channel}：本周已售 {row.current_week_qty:,.0f} 件，"
                f"预测整周 {row.forecast_week_qty:,.0f} 件，较上周 {row.week_qty_change:+.1%}。"
            )
    lines.extend(["", "## 近几周/日对比分析"])
    daily_trend = tables["daily_trend"]
    if daily_trend.empty:
        lines.append("- 暂无日趋势数据。")
    else:
        latest_days = sorted(daily_trend["date"].drop_duplicates())[-3:]
        day_summary = (
            daily_trend[daily_trend["date"].isin(latest_days)]
            .groupby("date", as_index=False)["sales_qty"]
            .sum()
            .sort_values("date")
        )
        for row in day_summary.itertuples(index=False):
            lines.append(f"- {row.date:%Y-%m-%d}：总销量 {row.sales_qty:,.0f} 件。")
    weekly_history = tables["weekly_history"]
    if not weekly_history.empty:
        week_summary = (
            weekly_history.groupby("week_start", as_index=False)["sales_qty"]
            .sum()
            .sort_values("week_start")
        )
        for row in week_summary.itertuples(index=False):
            lines.append(f"- {row.week_start:%Y-%m-%d} 周：累计销量 {row.sales_qty:,.0f} 件。")
    lines.extend(["", "## 头部商品"])
    sku = tables["sku"].head(10)
    if sku.empty:
        lines.append("- 暂无 SKU 数据。")
    else:
        for row in sku.itertuples(index=False):
            lines.append(f"- {row.channel}｜{row.sku_name}：销量 {row.sales_qty:,.0f} 件，库存 {row.stock_qty:,.0f} 件")
    lines.extend(["", "## 异常与行动"])
    anomalies = tables["anomalies"]
    if anomalies.empty:
        lines.append("- 当日未识别到明显异常。")
    else:
        for row in anomalies.head(10).itertuples(index=False):
            lines.append(f"- {row.channel}｜{row.sku_name}｜{row.issue}：{row.action}")
    lines.extend(["", "## 今日行动清单"])
    lines.extend(f"{index}. {item}" for index, item in enumerate(report["actions"], start=1))
    return "\n".join(lines)


def export_daily_report_excel(
    data: pd.DataFrame,
    report_date: pd.Timestamp,
    context: dict[str, object] | None = None,
) -> bytes:
    report = build_daily_report(data, report_date, context)
    tables = report["tables"]
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        summary_rows = [{"条目": "报告标题", "内容": report["title"]}]
        summary_rows.extend({"条目": f"经营结论 {index}", "内容": item} for index, item in enumerate(report["summary"], start=1))
        summary_rows.extend({"条目": f"行动建议 {index}", "内容": item} for index, item in enumerate(report["actions"], start=1))
        pd.DataFrame(summary_rows).to_excel(writer, sheet_name="经营摘要", index=False)
        pd.DataFrame([report["context"]]).to_excel(writer, sheet_name="日期天气活动", index=False)
        _chinese_table(report["analysis"]).to_excel(writer, sheet_name="涨跌原因", index=False)
        _chinese_table(tables["weekly"]).to_excel(writer, sheet_name="本周预测", index=False)
        _chinese_table(tables["weekly_history"]).to_excel(writer, sheet_name="近几周对比", index=False)
        _chinese_table(tables["channel"]).to_excel(writer, sheet_name="渠道对比", index=False)
        _chinese_table(tables["sku"].head(100)).to_excel(writer, sheet_name="SKU排行", index=False)
        _chinese_table(tables["daily_trend"]).to_excel(writer, sheet_name="日趋势", index=False)
        _chinese_table(tables["anomalies"]).to_excel(writer, sheet_name="异常行动", index=False)
        _action_register(tables["anomalies"], report_date).to_excel(writer, sheet_name="行动跟进", index=False)
        _write_leadership_summary(
            writer,
            report["title"],
            f"{report_date:%Y-%m-%d}",
            data[data["date"] <= report_date],
            generate_daily_report(data, report_date, context),
            _action_register(tables["anomalies"], report_date),
        )
        _style_workbook(writer)
    return output.getvalue()


REPORT_TYPE_LABELS = {"daily": "日报", "weekly": "周报", "monthly": "月报"}


def _day_end(value: pd.Timestamp) -> pd.Timestamp:
    return pd.Timestamp(value).normalize() + pd.Timedelta(days=1) - pd.Timedelta(nanoseconds=1)


def report_period(report_type: str, report_date: pd.Timestamp) -> tuple[pd.Timestamp, pd.Timestamp, pd.Timestamp, pd.Timestamp]:
    report_date = pd.Timestamp(report_date).normalize()
    if report_type == "daily":
        start = report_date
        end = _day_end(report_date)
        previous_start = report_date - pd.Timedelta(days=1)
        previous_end = _day_end(previous_start)
    elif report_type == "weekly":
        start = report_date - pd.Timedelta(days=report_date.weekday())
        end = _day_end(start + pd.Timedelta(days=6))
        previous_start = start - pd.Timedelta(days=7)
        previous_end = _day_end(previous_start + pd.Timedelta(days=6))
    elif report_type == "monthly":
        start = report_date.replace(day=1)
        end = _day_end(start + pd.offsets.MonthEnd(0))
        previous_month_end = start - pd.Timedelta(days=1)
        previous_start = previous_month_end.replace(day=1)
        previous_end = _day_end(previous_month_end)
    else:
        raise ValueError(f"不支持的报告类型: {report_type}")
    return start, end, previous_start, previous_end


def year_over_year_period(start: pd.Timestamp, end: pd.Timestamp) -> tuple[pd.Timestamp, pd.Timestamp]:
    return pd.Timestamp(start) - pd.DateOffset(years=1), pd.Timestamp(end) - pd.DateOffset(years=1)


def period_report_tables(data: pd.DataFrame, report_type: str, report_date: pd.Timestamp) -> dict[str, pd.DataFrame]:
    data = _with_purchase_column(data)
    start, end, previous_start, previous_end = report_period(report_type, report_date)
    yoy_start, yoy_end = year_over_year_period(start, end)
    current = data[(data["date"] >= start) & (data["date"] <= end)]
    previous = data[(data["date"] >= previous_start) & (data["date"] <= previous_end)]
    yoy = data[(data["date"] >= yoy_start) & (data["date"] <= yoy_end)]
    current_with_day = current.copy()
    if not current_with_day.empty:
        current_with_day["_report_day"] = pd.to_datetime(current_with_day["date"]).dt.normalize()

    channel = (
        current_with_day.groupby("channel", as_index=False)
        .agg(
            sales_qty=("sales_qty", "sum"),
            sales_amount=("sales_amount", "sum"),
            purchase_qty=("purchase_qty", "sum"),
            active_days=("_report_day", "nunique"),
        )
        if not current_with_day.empty else _empty_frame(["channel", "sales_qty", "sales_amount", "purchase_qty", "active_days"])
    )
    previous_channel = (
        previous.groupby("channel", as_index=False)
        .agg(previous_sales_qty=("sales_qty", "sum"), previous_sales_amount=("sales_amount", "sum"))
        if not previous.empty else _empty_frame(["channel", "previous_sales_qty", "previous_sales_amount"])
    )
    yoy_channel = (
        yoy.groupby("channel", as_index=False)
        .agg(yoy_sales_qty=("sales_qty", "sum"), yoy_sales_amount=("sales_amount", "sum"))
        if not yoy.empty else _empty_frame(["channel", "yoy_sales_qty", "yoy_sales_amount"])
    )
    channel = channel.merge(previous_channel, on="channel", how="outer").merge(yoy_channel, on="channel", how="outer")
    for column in (
        "sales_qty", "sales_amount", "purchase_qty", "active_days", "previous_sales_qty", "previous_sales_amount",
        "yoy_sales_qty", "yoy_sales_amount",
    ):
        if column in channel:
            channel[column] = pd.to_numeric(channel[column], errors="coerce").fillna(0)
    channel["qty_change"] = 0.0
    has_previous = channel["previous_sales_qty"] != 0
    channel.loc[has_previous, "qty_change"] = (
        (channel.loc[has_previous, "sales_qty"] - channel.loc[has_previous, "previous_sales_qty"])
        / channel.loc[has_previous, "previous_sales_qty"]
    )
    channel["yoy_qty_change"] = 0.0
    has_yoy = channel["yoy_sales_qty"] != 0
    channel.loc[has_yoy, "yoy_qty_change"] = (
        (channel.loc[has_yoy, "sales_qty"] - channel.loc[has_yoy, "yoy_sales_qty"])
        / channel.loc[has_yoy, "yoy_sales_qty"]
    )
    channel["purchase_sales_ratio"] = 0.0
    has_sales = channel["sales_qty"] != 0
    channel.loc[has_sales, "purchase_sales_ratio"] = channel.loc[has_sales, "purchase_qty"] / channel.loc[has_sales, "sales_qty"]

    sku_totals = (
        current.groupby(["channel", "sku_id", "sku_name"], as_index=False)
        .agg(sales_qty=("sales_qty", "sum"), sales_amount=("sales_amount", "sum"), purchase_qty=("purchase_qty", "sum"))
        .sort_values("sales_qty", ascending=False)
        if not current.empty else _empty_frame(["channel", "sku_id", "sku_name", "sales_qty", "sales_amount", "purchase_qty"])
    )
    daily = _daily_channel(current)
    latest = latest_inventory_rows(current)
    inventory = (
        latest.groupby(["channel", "sku_id", "sku_name"], as_index=False)
        .agg(stock_qty=("stock_qty", "sum"), latest_sales_qty=("sales_qty", "sum"))
        .sort_values("stock_qty")
        if not latest.empty else _empty_frame(["channel", "sku_id", "sku_name", "stock_qty", "latest_sales_qty"])
    )
    sku = (
        sku_totals.merge(
            inventory[["channel", "sku_id", "stock_qty"]],
            on=["channel", "sku_id"],
            how="left",
        ).fillna({"stock_qty": 0})
        if not sku_totals.empty
        else _empty_frame(["channel", "sku_id", "sku_name", "sales_qty", "sales_amount", "purchase_qty", "stock_qty"])
    )
    if not sku.empty:
        sku["purchase_sales_gap"] = sku["purchase_qty"] - sku["sales_qty"]
    covered_days = max(int(pd.to_datetime(current["date"]).dt.normalize().nunique()) if not current.empty else 0, 1)
    if report_type == "weekly":
        full_days = 7
    elif report_type == "monthly":
        full_days = int(end.days_in_month)
    else:
        full_days = 1
    forecast = channel[["channel", "sales_qty", "previous_sales_qty", "qty_change"]].copy()
    forecast["forecast_qty"] = forecast["sales_qty"] / covered_days * full_days
    forecast["remaining_days"] = max(full_days - covered_days, 0)
    anomalies = detect_anomalies(data[data["date"] <= end])
    actions = _action_register(anomalies, end)
    return {
        "channel": channel.sort_values("sales_qty", ascending=False),
        "sku": sku,
        "daily": daily,
        "inventory": inventory,
        "forecast": forecast.sort_values("forecast_qty", ascending=False),
        "anomalies": anomalies,
        "actions": actions,
    }


def generate_period_report(
    data: pd.DataFrame,
    report_type: str,
    report_date: pd.Timestamp,
    context: dict[str, object] | None = None,
) -> str:
    data = _with_purchase_column(data)
    label = REPORT_TYPE_LABELS[report_type]
    start, end, previous_start, previous_end = report_period(report_type, report_date)
    yoy_start, yoy_end = year_over_year_period(start, end)
    tables = period_report_tables(data, report_type, report_date)
    current = data[(data["date"] >= start) & (data["date"] <= end)]
    previous = data[(data["date"] >= previous_start) & (data["date"] <= previous_end)]
    yoy = data[(data["date"] >= yoy_start) & (data["date"] <= yoy_end)]
    context_rows = calendar_context(end, context)
    if current.empty:
        return f"# 渠道销售{label} | {start:%Y-%m-%d} 至 {end:%Y-%m-%d}\n\n当前周期无数据。"

    qty = current["sales_qty"].sum()
    previous_qty = previous["sales_qty"].sum()
    qty_change = (qty - previous_qty) / previous_qty if previous_qty else 0
    yoy_qty = yoy["sales_qty"].sum()
    yoy_change = (qty - yoy_qty) / yoy_qty if yoy_qty else 0
    purchase_qty = current["purchase_qty"].sum()
    purchase_ratio = purchase_qty / qty if qty else 0
    purchase_text = (
        f"采购数量 {purchase_qty:,.0f} 件，采销比 {purchase_ratio:.1%}"
        if purchase_qty
        else "本期未导入采购订单或未识别采购数量"
    )
    amount = current["sales_amount"].sum()
    amount_channels = current.groupby("channel")["sales_amount"].apply(lambda values: values.abs().sum() > 0)
    amount_text = (
        f"销售额 ¥{amount:,.0f}" if amount > 0 and amount_channels.all()
        else f"已提供渠道销售额 ¥{amount:,.0f}，未覆盖全部渠道" if amount > 0
        else "销售额字段不完整"
    )
    top_channel = tables["channel"].iloc[0]["channel"] if not tables["channel"].empty else "无"
    forecast_qty = tables["forecast"]["forecast_qty"].sum() if not tables["forecast"].empty else qty
    covered_days = int(current["date"].nunique())
    expected_days = int((end - start).days + 1)
    coverage_text = (
        f"数据覆盖完整，共 {covered_days} 天。"
        if covered_days >= expected_days
        else f"周期应有 {expected_days} 天，当前数据覆盖 {covered_days} 天，缺少 {expected_days - covered_days} 天；预测按已覆盖日期日均推算。"
    )
    comparison_text = (
        f"较上一同期 {'增长' if qty_change > 0 else '下降' if qty_change < 0 else '持平'} {abs(qty_change):.1%}"
        if previous_qty
        else "上一同期无可比销量数据"
    )
    yoy_text = (
        f"同比去年同期 {'增长' if yoy_change > 0 else '下降' if yoy_change < 0 else '持平'} {abs(yoy_change):.1%}"
        if yoy_qty
        else "同比数据不足"
    )
    comparison_label = {"daily": "上一日", "weekly": "上一周", "monthly": "上月"}[report_type]
    lines = [
        f"# 渠道销售{label} | {start:%Y-%m-%d} 至 {end:%Y-%m-%d}",
        "",
        "## 一、管理层摘要",
        f"- 本期销量 {qty:,.0f} 件，{comparison_text}，{yoy_text}；{amount_text}；{purchase_text}。",
        f"- {top_channel} 为本期销量最高渠道；按当前日均节奏预测完整周期销量 {forecast_qty:,.0f} 件。",
        f"- 日期/节假日：{context_rows['日期标签']}、{context_rows['节假日']}；天气：{context_rows['天气']}；渠道活动：{context_rows['渠道活动']}。",
        "",
        "## 二、核心指标与上期对比",
        f"- 当前周期：{start:%Y-%m-%d} 至 {end:%Y-%m-%d}；上一同期：{previous_start:%Y-%m-%d} 至 {previous_end:%Y-%m-%d}；去年同期：{yoy_start:%Y-%m-%d} 至 {yoy_end:%Y-%m-%d}。",
        f"- 覆盖渠道 {current['channel'].nunique()} 个，SKU {current['sku_id'].nunique()} 个，销量 {qty:,.0f} 件，采购 {purchase_qty:,.0f} 件，上一同期 {previous_qty:,.0f} 件，去年同期 {yoy_qty:,.0f} 件；{yoy_text}。",
        f"- 数据完整性：{coverage_text}",
        "",
        "## 三、销售 Pipeline 分析",
    ]
    lines.extend(_pipeline_markdown_lines(context, end))
    lines.extend(["", "## 四、渠道表现与涨跌驱动"])
    for row in tables["channel"].itertuples(index=False):
        row_series = pd.Series(row._asdict())
        reason = _reason_for_channel(row_series, context_rows, comparison_label)
        channel_comparison = (
            f"较{comparison_label} {row.qty_change:+.1%}"
            if row.previous_sales_qty
            else f"{comparison_label}无可比数据"
        )
        channel_yoy = f"同比 {row.yoy_qty_change:+.1%}" if row.yoy_sales_qty else "同比数据不足"
        purchase_clause = f"，采购 {row.purchase_qty:,.0f} 件，采销比 {row.purchase_sales_ratio:.1%}" if row.purchase_qty else "，未导入采购数量"
        lines.append(
            f"- **{row.channel}**：销量 {row.sales_qty:,.0f} 件{purchase_clause}，"
            f"{channel_comparison}，{channel_yoy}。{reason}"
        )
    lines.extend(["", "## 五、SKU 贡献与机会"])
    for row in tables["sku"].head(10).itertuples(index=False):
        gap = getattr(row, "purchase_sales_gap", 0)
        lines.append(
            f"- {row.channel}｜{row.sku_name}：销量 {row.sales_qty:,.0f} 件，"
            f"采购 {row.purchase_qty:,.0f} 件，采销差 {gap:,.0f} 件，库存 {row.stock_qty:,.0f} 件。"
        )
    lines.extend(["", "## 六、采销匹配分析"])
    if purchase_qty:
        gap_table = tables["sku"].copy()
        gap_table["abs_gap"] = gap_table["purchase_sales_gap"].abs()
        for row in gap_table.sort_values("abs_gap", ascending=False).head(6).itertuples(index=False):
            direction = "采购高于销售" if row.purchase_sales_gap > 0 else "采购低于销售" if row.purchase_sales_gap < 0 else "采销平衡"
            lines.append(
                f"- {row.channel}｜{row.sku_name}：{direction} {abs(row.purchase_sales_gap):,.0f} 件，"
                f"销售 {row.sales_qty:,.0f} 件，采购 {row.purchase_qty:,.0f} 件。"
            )
    else:
        lines.append("- 本期没有可用采购订单数据，暂不输出采销匹配判断。")
    lines.extend(["", "## 七、库存风险与异常"])
    if tables["anomalies"].empty:
        lines.append("- 当前未识别到明显库存或销量异常。")
    else:
        for row in tables["anomalies"].head(10).itertuples(index=False):
            lines.append(f"- {row.channel}｜{row.sku_name}｜{row.issue}：{row.action}")
    lines.extend(["", "## 八、销量预测"])
    for row in tables["forecast"].itertuples(index=False):
        lines.append(f"- {row.channel}：当前累计 {row.sales_qty:,.0f} 件，完整周期预测 {row.forecast_qty:,.0f} 件。")
    lines.extend([
        "",
        "## 九、行动跟进表",
    ])
    if tables["actions"].empty:
        lines.append("- 当前没有需要跟进的异常事项。")
    else:
        for row in tables["actions"].head(10).itertuples(index=False):
            lines.append(
                f"- **{row.优先级}｜{row.渠道}｜{row.SKU}**：{row.问题}；{row.建议动作}；"
                f"负责人：{row.负责人}；截止：{row.截止日期}；状态：{row.处理状态}。"
            )
    lines.extend([
        "",
        "## 十、外部因素与待验证假设",
        f"- 天气：{context_rows['天气']}",
        f"- 节假日/特殊日期：{context_rows['节假日']}",
        f"- 渠道活动/价格/上架变化：{context_rows['渠道活动']}",
        "",
        "## 十一、复盘要求",
        "1. 业务负责人填写行动跟进表中的负责人、处理状态和实际结果。",
        "2. 下一周期复核 P0/P1 事项是否改善销量、库存和履约。",
        "3. 补充缺失的销售额、天气和活动数据，提升下一周期归因准确度。",
    ])
    return "\n".join(lines)


def _format_int(value: object) -> str:
    return f"{float(value or 0):,.0f}"


def _format_currency(value: object, has_amount: bool = True) -> str:
    if not has_amount:
        return "未提供"
    return f"¥{float(value or 0):,.0f}"


def _format_change(current: float, baseline: float) -> str:
    if not baseline:
        return "无基线"
    change = (current - baseline) / baseline
    if change > 0:
        return f"增长 {abs(change):.1%}"
    if change < 0:
        return f"下降 {abs(change):.1%}"
    return "持平 0.0%"


def _trim_text(value: object, limit: int = 34) -> str:
    text = str(value or "").strip()
    return text if len(text) <= limit else f"{text[:limit - 1]}…"


def _set_docx_font(run, name: str = "Arial", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.font.bold = bold


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line_spacing: float = 1.1) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell, value: object, bold: bool = False, color: str = "1D1D1F", fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _set_paragraph_spacing(paragraph, after=0, line_spacing=1.05)
    run = paragraph.add_run(str(value))
    _set_docx_font(run, size=9.5, color=color, bold=bold)
    if fill:
        _shade_cell(cell, fill)


def _add_section_heading(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, before=10, after=4, line_spacing=1.1)
    run = paragraph.add_run(text)
    _set_docx_font(run, size=13, color="1D1D1F", bold=True)


def _add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    _set_paragraph_spacing(paragraph, after=4, line_spacing=1.12)
    run = paragraph.add_run(text)
    _set_docx_font(run, size=10.5, color="30343B")


def _add_table(doc, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True, color="1D1D1F", fill="F2F4F7")
    if not rows:
        row = table.add_row().cells
        row[0].merge(row[-1])
        _set_cell_text(row[0], "暂无数据", color="6E6E73")
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            _set_cell_text(cell, value)
    doc.add_paragraph()


def export_period_report_docx(
    data: pd.DataFrame,
    report_type: str,
    report_date: pd.Timestamp,
    context: dict[str, object] | None = None,
) -> bytes:
    """Export a concise management-ready Word report."""
    data = _with_purchase_column(data)
    from docx import Document
    from docx.shared import Inches, Pt

    label = REPORT_TYPE_LABELS[report_type]
    start, end, previous_start, previous_end = report_period(report_type, report_date)
    yoy_start, yoy_end = year_over_year_period(start, end)
    tables = period_report_tables(data, report_type, report_date)
    current = data[(data["date"] >= start) & (data["date"] <= end)]
    previous = data[(data["date"] >= previous_start) & (data["date"] <= previous_end)]
    yoy = data[(data["date"] >= yoy_start) & (data["date"] <= yoy_end)]
    context_rows = calendar_context(end, context)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = None

    title = doc.add_paragraph()
    _set_paragraph_spacing(title, after=2, line_spacing=1.05)
    title_run = title.add_run(f"渠道销售{label}")
    _set_docx_font(title_run, size=22, color="1D1D1F", bold=True)

    subtitle = doc.add_paragraph()
    _set_paragraph_spacing(subtitle, after=10, line_spacing=1.1)
    subtitle_run = subtitle.add_run(f"{start:%Y-%m-%d} 至 {end:%Y-%m-%d} | 自动生成经营简报")
    _set_docx_font(subtitle_run, size=10, color="6E6E73")

    if current.empty:
        _add_bullet(doc, "当前周期无数据，无法生成经营判断。")
        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    qty = float(current["sales_qty"].sum())
    previous_qty = float(previous["sales_qty"].sum())
    yoy_qty = float(yoy["sales_qty"].sum())
    purchase_qty = float(current["purchase_qty"].sum())
    purchase_ratio = purchase_qty / qty if qty else 0.0
    amount = float(current["sales_amount"].sum())
    has_amount = bool(amount > 0 and current.groupby("channel")["sales_amount"].apply(lambda values: values.abs().sum() > 0).all())
    covered_days = int(current["date"].nunique())
    expected_days = int((end - start).days + 1)
    latest_stock = float(latest_inventory_rows(current)["stock_qty"].sum())
    top_channel = tables["channel"].iloc[0]["channel"] if not tables["channel"].empty else "无"

    kpi_rows = [
        ["本期销量", f"{_format_int(qty)} 件", "环比", _format_change(qty, previous_qty)],
        ["同比销量", f"{_format_int(yoy_qty)} 件" if yoy_qty else "未提供", "同比", _format_change(qty, yoy_qty)],
        ["本期采购量", f"{_format_int(purchase_qty)} 件" if purchase_qty else "未提供", "采销比", f"{purchase_ratio:.1%}" if purchase_qty else "无采购数据"],
        ["本期销售额", _format_currency(amount, has_amount), "最新库存", f"{_format_int(latest_stock)} 件"],
        ["覆盖范围", f"{current['channel'].nunique()} 渠道 / {current['sku_id'].nunique()} SKU", "数据覆盖", f"{covered_days}/{expected_days} 天"],
    ]
    _add_table(doc, ["指标", "结果", "对比", "说明"], kpi_rows)

    _add_section_heading(doc, "一、管理层摘要")
    coverage_note = "数据覆盖完整。" if covered_days >= expected_days else f"周期应有 {expected_days} 天，当前仅覆盖 {covered_days} 天。"
    amount_note = "销售额口径完整。" if has_amount else "销售额字段未覆盖全部渠道，金额判断仅作参考。"
    summary_items = [
        f"本期销量 {_format_int(qty)} 件，环比{_format_change(qty, previous_qty)}，同比{_format_change(qty, yoy_qty)}。",
        (
            f"本期采购数量 {_format_int(purchase_qty)} 件，采销比 {purchase_ratio:.1%}，"
            f"采购与销售差额 {_format_int(purchase_qty - qty)} 件。"
            if purchase_qty else "本期未导入采购订单或未识别采购数量，采销匹配判断暂不输出。"
        ),
        f"{top_channel} 为销量最高渠道；{amount_note}",
        f"{coverage_note} 天气：{context_rows['天气']}；节假日/特殊日期：{context_rows['节假日']}。",
    ]
    for item in summary_items:
        _add_bullet(doc, item)

    _add_section_heading(doc, "二、销售 Pipeline")
    pipeline_summary = _pipeline_summary(context, end)
    pipeline_open = pipeline_summary["open_frame"]
    if pipeline_open.empty:
        _add_bullet(doc, "当前未维护销售 Pipeline，建议补充客户公司、阶段、预计金额、成交概率和下一步动作。")
    else:
        pipeline_kpis = [
            ["开放商机", f"{pipeline_summary['open_count']:,} 个", "预计订单金额", f"¥{pipeline_summary['open_amount']:,.0f}"],
            ["加权预测金额", f"¥{pipeline_summary['weighted_amount']:,.0f}", "金额集中阶段", str(pipeline_summary["top_stage"])],
            ["30天内预计成交", f"{pipeline_summary['upcoming_count']:,} 个", "高风险项", f"{pipeline_summary['risk_count']:,} 条"],
        ]
        _add_table(doc, ["指标", "结果", "指标", "结果"], pipeline_kpis)
        pipeline_rows = []
        for _, row in (
            pipeline_open.sort_values(["weighted_amount", "Deal Size"], ascending=False)
            .head(6)
            .iterrows()
        ):
            close_date = row.get("Expected Close Date")
            close_text = f"{pd.Timestamp(close_date):%Y-%m-%d}" if pd.notna(close_date) else "未定"
            pipeline_rows.append([
                _trim_text(str(row.get("Account / Company", "")), 18),
                str(row.get("Stage", "")),
                f"¥{float(row.get('Deal Size', 0)):,.0f}",
                f"{float(row.get('Probability', 0)):.0%}",
                close_text,
                _trim_text(str(row.get("Next Step", "")), 28),
                _trim_text(str(row.get("Risk", "")), 22),
            ])
        _add_table(doc, ["客户", "阶段", "预计金额", "概率", "预计成交", "下一步", "风险"], pipeline_rows)

    _add_section_heading(doc, "三、渠道表现")
    channel_rows = []
    for row in tables["channel"].head(8).itertuples(index=False):
        channel_rows.append([
            row.channel,
            f"{_format_int(row.sales_qty)} 件",
            f"{_format_int(row.purchase_qty)} 件",
            f"{row.purchase_sales_ratio:.1%}" if row.purchase_qty else "无采购",
            _format_currency(row.sales_amount, row.sales_amount > 0),
            _format_change(float(row.sales_qty), float(row.previous_sales_qty)),
            _format_change(float(row.sales_qty), float(row.yoy_sales_qty)),
            f"{int(row.active_days)} 天",
        ])
    _add_table(doc, ["渠道", "销量", "采购量", "采销比", "销售额", "环比", "同比", "有效天数"], channel_rows)

    _add_section_heading(doc, "四、TOP SKU")
    sku_rows = []
    for row in tables["sku"].head(8).itertuples(index=False):
        sku_rows.append([
            row.channel,
            _trim_text(row.sku_name, 30),
            f"{_format_int(row.sales_qty)} 件",
            f"{_format_int(row.purchase_qty)} 件",
            f"{_format_int(row.purchase_sales_gap)} 件",
            _format_currency(row.sales_amount, row.sales_amount > 0),
            f"{_format_int(row.stock_qty)} 件",
        ])
    _add_table(doc, ["渠道", "商品", "销量", "采购量", "采销差", "销售额", "最新库存"], sku_rows)

    _add_section_heading(doc, "五、采销匹配")
    if purchase_qty:
        gap_rows = []
        gap_table = tables["sku"].copy()
        gap_table["abs_gap"] = gap_table["purchase_sales_gap"].abs()
        for row in gap_table.sort_values("abs_gap", ascending=False).head(8).itertuples(index=False):
            direction = "采购高于销售" if row.purchase_sales_gap > 0 else "采购低于销售" if row.purchase_sales_gap < 0 else "采销平衡"
            gap_rows.append([
                row.channel,
                _trim_text(row.sku_name, 28),
                f"{_format_int(row.sales_qty)} 件",
                f"{_format_int(row.purchase_qty)} 件",
                direction,
                f"{_format_int(abs(row.purchase_sales_gap))} 件",
            ])
        _add_table(doc, ["渠道", "商品", "销量", "采购量", "判断", "差额"], gap_rows)
    else:
        _add_bullet(doc, "本期没有可用采购订单数据，暂不输出采销匹配判断。")

    _add_section_heading(doc, "六、行动跟进")
    action_rows = []
    for row in tables["actions"].head(8).itertuples(index=False):
        action_rows.append([
            row.优先级,
            row.渠道,
            _trim_text(row.SKU, 22),
            row.问题,
            _trim_text(row.建议动作, 34),
            row.截止日期,
        ])
    _add_table(doc, ["优先级", "渠道", "商品", "问题", "建议动作", "截止"], action_rows)

    _add_section_heading(doc, "七、口径说明")
    notes = [
        f"当前周期：{start:%Y-%m-%d} 至 {end:%Y-%m-%d}；上一同期：{previous_start:%Y-%m-%d} 至 {previous_end:%Y-%m-%d}；去年同期：{yoy_start:%Y-%m-%d} 至 {yoy_end:%Y-%m-%d}。",
        "库存使用渠道 + SKU 的最新快照口径，不跨日期累加。",
        "采购数量来自识别为采购订单的数据，独立于销量口径；采购金额不会计入销售额。",
        f"渠道活动/价格/上架变化：{context_rows['渠道活动']}。",
    ]
    for item in notes:
        _add_bullet(doc, item)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def export_period_report_excel(
    data: pd.DataFrame,
    report_type: str,
    report_date: pd.Timestamp,
    context: dict[str, object] | None = None,
) -> bytes:
    label = REPORT_TYPE_LABELS[report_type]
    start, end, previous_start, previous_end = report_period(report_type, report_date)
    yoy_start, yoy_end = year_over_year_period(start, end)
    tables = period_report_tables(data, report_type, report_date)
    report_text = generate_period_report(data, report_type, report_date, context)
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame({
            "项目": ["报告类型", "当前周期", "上一同期", "去年同期", "天气", "节假日", "渠道活动"],
            "内容": [
                label, f"{start:%Y-%m-%d} 至 {end:%Y-%m-%d}", f"{previous_start:%Y-%m-%d} 至 {previous_end:%Y-%m-%d}",
                f"{yoy_start:%Y-%m-%d} 至 {yoy_end:%Y-%m-%d}",
                _context_value(context, "weather_note"), _context_value(context, "holiday_note"), _context_value(context, "activity_note"),
            ],
        }).to_excel(writer, sheet_name="报告说明", index=False)
        pd.DataFrame({"报告正文": report_text.splitlines()}).to_excel(writer, sheet_name="经营结论", index=False)
        _chinese_table(tables["channel"]).to_excel(writer, sheet_name="渠道对比", index=False)
        _chinese_table(tables["sku"].head(200)).to_excel(writer, sheet_name="SKU贡献", index=False)
        _chinese_table(tables["daily"]).to_excel(writer, sheet_name="每日趋势", index=False)
        _chinese_table(tables["inventory"]).to_excel(writer, sheet_name="库存风险", index=False)
        _chinese_table(tables["forecast"]).to_excel(writer, sheet_name="销量预测", index=False)
        _chinese_table(tables["anomalies"]).to_excel(writer, sheet_name="异常行动", index=False)
        tables["actions"].to_excel(writer, sheet_name="行动跟进", index=False)
        current = data[(data["date"] >= start) & (data["date"] <= end)]
        _write_leadership_summary(
            writer,
            f"渠道销售{label}",
            f"{start:%Y-%m-%d} 至 {end:%Y-%m-%d}",
            current,
            report_text,
            tables["actions"],
        )
        _style_workbook(writer)
    return output.getvalue()
