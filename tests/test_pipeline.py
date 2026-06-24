from pathlib import Path
import sys
from io import BytesIO

import pandas as pd
import pytest
from openpyxl import load_workbook

sys.path.insert(0, str(Path(__file__).parents[1] / "src"))

from retail_assistant.analytics import (
    calculate_kpis,
    channel_summary,
    detect_anomalies,
    generate_management_summary,
    latest_inventory_rows,
    sku_summary,
)
from retail_assistant.charts import choose_available_metric, horizontal_bar_chart, trend_chart
from retail_assistant.normalize import (
    combine_sales_files,
    infer_channel,
    infer_mapping,
    normalize_sales_data,
    profile_file,
    read_sales_file,
)
from retail_assistant.reporting import (
    daily_report_tables,
    export_daily_report_excel,
    export_period_report_excel,
    export_period_report_docx,
    generate_daily_report,
    generate_period_report,
    period_report_tables,
    report_period,
)
from retail_assistant.weather import fetch_regional_weather, format_weather_note


def test_normalize_maps_chinese_columns_and_only_drops_exact_duplicates():
    raw = pd.DataFrame(
        {
            "日期": ["2026-06-10", "2026-06-10", "2026-06-10"],
            "SKU编码": ["A1", "A1", "A1"],
            "SKU名称": ["测试商品", "测试商品", "测试商品"],
            "销量": [8, 8, 9],
            "销售额": [80, 80, 90],
        }
    )

    result = normalize_sales_data(raw, default_channel="测试渠道")

    assert len(result) == 2
    assert result["sales_qty"].sum() == 17


def test_semantic_mapping_handles_unseen_but_descriptive_headers():
    mapping = infer_mapping(pd.Index(["统计时间", "货品名称", "货品ID", "本日销售量", "成交金额", "剩余库存"]))

    assert mapping["date"] == "统计时间"
    assert mapping["sku_name"] == "货品名称"
    assert mapping["sku_id"] == "货品ID"
    assert mapping["sales_qty"] == "本日销售量"
    assert mapping["sales_amount"] == "成交金额"
    assert mapping["stock_qty"] == "剩余库存"


def test_read_short_csv_does_not_fail_during_header_detection(tmp_path):
    path = tmp_path / "short.csv"
    path.write_text("日期,商品ID,商品名称,销量\n2026-06-10,A1,商品A,2\n", encoding="utf-8")

    result = read_sales_file(path)

    assert list(result.columns) == ["日期", "商品ID", "商品名称", "销量"]


def test_read_short_excel_and_stock_only_file_can_render(tmp_path):
    path = tmp_path / "short-stock.xlsx"
    pd.DataFrame(
        {
            "业务日期": [20260610, 20260610],
            "商品编码": ["H1", "H2"],
            "商品名称": ["商品一", "商品二"],
            "经营店名称": ["上海浦东店", "苏州园区店"],
            "库存总数": [120, 80],
        }
    ).to_excel(path, index=False)

    raw = read_sales_file(path)
    profile = profile_file(raw, path.name)
    result = combine_sales_files([{
        "frame": raw,
        "channel": "盒马",
        "source_file": path.name,
        "data_type": "库存",
        "mapping": profile["mapping"],
    }])
    metric, note = choose_available_metric(result, "sales_qty")

    assert len(result) == 2
    assert result["sales_qty"].sum() == 0
    assert result["stock_qty"].sum() == 200
    assert metric == "stock_qty" and "自动改为展示库存" in note


def test_profile_identifies_real_channel_signatures():
    hema = pd.DataFrame(columns=["业务日期", "经营店名称", "供应商编码", "商品编码", "销售数量"])
    xiaoxiang = pd.DataFrame(columns=["日期", "商品ID", "商品合作模式", "前置站点库存数量", "商品销售量"])

    assert infer_channel(hema, "5.8-6.9") == "盒马"
    assert infer_channel(xiaoxiang, "商品明细") == "小象"
    assert profile_file(hema, "5.8-6.9.xlsx")["data_type"] == "销售"


def test_normalize_real_hema_columns_and_compact_date():
    raw = pd.DataFrame(
        {
            "业务日期": [20260609],
            "商品编码": [441588014],
            "商品名称": ["测试葡萄酒"],
            "子公司名称": ["上海子公司"],
            "经营店名称": ["上海门店"],
            "销售数量": [34],
            "库存总数": [152],
        }
    )

    result = normalize_sales_data(raw, default_channel="盒马", source_file="盒马.xlsx")

    assert result.loc[0, "date"] == pd.Timestamp("2026-06-09")
    assert result.loc[0, "sales_amount"] == 0
    assert result.loc[0, "stock_qty"] == 152
    assert result.loc[0, "store"] == "上海门店"
    assert result.loc[0, "source_file"] == "盒马.xlsx"


def test_hema_sales_only_keep_operating_store_names_containing_store_character():
    raw = pd.DataFrame(
        {
            "业务日期": [20260609] * 4,
            "商品编码": ["H1"] * 4,
            "商品名称": ["盒马商品"] * 4,
            "经营店名称": ["上海浦东店", "华东FDC", "上海前置仓", "北京鲜生会员店"],
            "销售数量": [10, 100, 200, 20],
        }
    )

    result = normalize_sales_data(raw, default_channel="盒马", source_file="盒马销售.xlsx", data_type="销售")

    assert result["sales_qty"].sum() == 30
    assert set(result["store"]) == {"上海浦东店", "北京鲜生会员店"}
    assert "已排除 2 条前置仓、FDC 或其他非门店记录" in result.attrs["warnings"][-1]


def test_combine_preserves_channels_skus_dates_and_source_files():
    hema = pd.DataFrame(
        {
            "业务日期": [20260609, 20260610],
            "商品编码": ["H1", "H1"],
            "商品名称": ["盒马商品", "盒马商品"],
            "销售数量": [10, 12],
            "库存总数": [30, 20],
        }
    )
    xiaoxiang = pd.DataFrame(
        {
            "日期": [20260609, 20260610],
            "商品ID": ["X1", "X1"],
            "商品名称": ["小象商品", "小象商品"],
            "商品销售量": [8, 9],
            "商品销售额": [80, 90],
            "前置站点库存数量": [20, 10],
        }
    )
    items = [
        {"frame": hema, "channel": "盒马", "source_file": "hema.xlsx", "data_type": "销售库存"},
        {"frame": xiaoxiang, "channel": "小象", "source_file": "xiaoxiang.xlsx", "data_type": "销售库存"},
    ]

    result = combine_sales_files(items)

    assert set(result["channel"]) == {"盒马", "小象"}
    assert set(result["sku_id"]) == {"H1", "X1"}
    assert result["date"].nunique() == 2
    assert set(result["source_file"]) == {"hema.xlsx", "xiaoxiang.xlsx"}


def test_combine_removes_overlapping_business_rows_across_files():
    raw = pd.DataFrame(
        {
            "日期": [20260610],
            "商品ID": ["X1"],
            "商品名称": ["小象商品"],
            "城市": ["上海"],
            "商品销售量": [9],
            "商品销售额": [90],
            "前置站点库存数量": [10],
        }
    )
    result = combine_sales_files([
        {"frame": raw, "channel": "小象", "source_file": "周报.xlsx", "data_type": "销售库存"},
        {"frame": raw, "channel": "小象", "source_file": "月报.xlsx", "data_type": "销售库存"},
    ])

    assert len(result) == 1
    assert "移除了 1 条重复业务记录" in result.attrs["warnings"][-1]


def test_anomalies_are_aggregated_by_channel_and_sku_not_store_rows():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-09", "2026-06-09", "2026-06-10", "2026-06-10"]),
            "channel": ["盒马"] * 4,
            "sku_id": ["A1"] * 4,
            "sku_name": ["测试商品"] * 4,
            "city": ["上海"] * 4,
            "store": ["店1", "店2", "店1", "店2"],
            "sales_qty": [5, 5, 7, 5],
            "sales_amount": [50, 50, 70, 50],
            "stock_qty": [10, 10, 0, 0],
        }
    )

    anomalies = detect_anomalies(data)

    assert len(anomalies) == 1
    assert anomalies.iloc[0]["sales_qty"] == 12
    assert anomalies.iloc[0]["issue"] == "缺货风险"


def test_channel_summary_and_kpis_keep_channels_separate():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-10", "2026-06-10"]),
            "channel": ["盒马", "小象"],
            "sku_id": ["A1", "B1"],
            "sku_name": ["A", "B"],
            "city": ["上海", "北京"],
            "sales_qty": [10, 20],
            "sales_amount": [0, 300],
            "stock_qty": [30, 40],
        }
    )

    assert calculate_kpis(data)["channel_count"] == 2
    summary = channel_summary(data)
    assert summary.iloc[0]["channel"] == "小象"
    assert summary["qty_share"].sum() == pytest.approx(1)
    management_summary = generate_management_summary(data, pd.DataFrame())
    assert "已提供渠道销售额" in management_summary


def test_latest_inventory_uses_each_channel_sku_latest_snapshot_not_all_dates():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-09", "2026-06-10", "2026-06-09"]),
            "channel": ["小象", "小象", "盒马"],
            "sku_id": ["X1", "X1", "H1"],
            "sku_name": ["小象商品", "小象商品", "盒马商品"],
            "city": ["上海", "上海", "上海"],
            "sales_qty": [1, 2, 3],
            "sales_amount": [10, 20, 0],
            "stock_qty": [100, 80, 30],
        }
    )

    latest = latest_inventory_rows(data)

    assert latest["stock_qty"].sum() == 110
    assert calculate_kpis(data)["latest_stock"] == 110


def test_empty_sku_summary_keeps_sortable_schema():
    summary = sku_summary(pd.DataFrame())

    assert summary.empty
    assert {"sales_qty", "sku_id", "sku_name", "latest_stock"} <= set(summary.columns)
    assert summary.sort_values("sales_qty").empty


def test_daily_report_contains_channel_and_sku_sections():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-09", "2026-06-10", "2026-06-09", "2026-06-10"]),
            "channel": ["盒马", "盒马", "小象", "小象"],
            "sku_id": ["H1", "H1", "X1", "X1"],
            "sku_name": ["盒马商品", "盒马商品", "小象商品", "小象商品"],
            "city": ["上海", "上海", "北京", "北京"],
            "sales_qty": [10, 12, 20, 25],
            "sales_amount": [0, 0, 200, 250],
            "stock_qty": [30, 20, 40, 30],
        }
    )

    report = generate_daily_report(data, pd.Timestamp("2026-06-10"))
    tables = daily_report_tables(data, pd.Timestamp("2026-06-10"))

    assert "# 渠道销售日报 | 2026-06-10" in report
    assert "盒马" in report and "小象" in report
    assert "已提供渠道销售额" in report
    assert "头部商品" in report and "今日行动清单" in report
    assert set(tables["channel"]["channel"]) == {"盒马", "小象"}


def test_daily_report_includes_context_forecast_reasons_and_excel_export():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime([
                "2026-06-01", "2026-06-02", "2026-06-08", "2026-06-09", "2026-06-10",
                "2026-06-01", "2026-06-02", "2026-06-08", "2026-06-09", "2026-06-10",
            ]),
            "channel": ["盒马"] * 5 + ["小象"] * 5,
            "sku_id": ["H1"] * 5 + ["X1"] * 5,
            "sku_name": ["盒马商品"] * 5 + ["小象商品"] * 5,
            "city": ["上海"] * 10,
            "sales_qty": [10, 10, 12, 13, 18, 20, 21, 20, 19, 16],
            "sales_amount": [100, 100, 120, 130, 180, 200, 210, 200, 190, 160],
            "stock_qty": [30, 30, 28, 26, 24, 40, 38, 36, 34, 32],
        }
    )
    context = {
        "weather_note": "上海高温，冷饮类可能受益",
        "holiday_note": "618 预热期",
        "activity_note": "盒马满减，小象无活动",
    }

    report = generate_daily_report(data, pd.Timestamp("2026-06-10"), context)
    tables = daily_report_tables(data, pd.Timestamp("2026-06-10"))
    workbook_bytes = export_daily_report_excel(data, pd.Timestamp("2026-06-10"), context)
    workbook = load_workbook(BytesIO(workbook_bytes), read_only=True)

    assert "日期、天气与活动背景" in report
    assert "渠道涨跌原因" in report
    assert "本周销量预测" in report
    assert "上海高温" in report and "盒马满减" in report
    assert not tables["weekly"].empty
    assert {"领导摘要", "经营摘要", "日期天气活动", "涨跌原因", "本周预测", "渠道对比", "SKU排行", "日趋势", "异常行动", "行动跟进"}.issubset(
        set(workbook.sheetnames)
    )
    assert workbook_bytes[:2] == b"PK"


@pytest.mark.parametrize(
    ("report_type", "label"),
    [("daily", "日报"), ("weekly", "周报"), ("monthly", "月报")],
)
def test_period_reports_support_daily_weekly_monthly(report_type, label):
    dates = pd.date_range("2026-05-01", "2026-06-10", freq="D")
    data = pd.DataFrame(
        {
            "date": list(dates) * 2,
            "channel": ["盒马"] * len(dates) + ["小象"] * len(dates),
            "sku_id": ["H1"] * len(dates) + ["X1"] * len(dates),
            "sku_name": ["盒马商品"] * len(dates) + ["小象商品"] * len(dates),
            "city": ["上海"] * (len(dates) * 2),
            "sales_qty": [10] * len(dates) + [20] * len(dates),
            "sales_amount": [0] * len(dates) + [200] * len(dates),
            "stock_qty": [30] * (len(dates) * 2),
        }
    )
    context = {"weather_note": "高温", "holiday_note": "618预热", "activity_note": "渠道满减"}

    report = generate_period_report(data, report_type, pd.Timestamp("2026-06-10"), context)
    tables = period_report_tables(data, report_type, pd.Timestamp("2026-06-10"))
    workbook = load_workbook(
        BytesIO(export_period_report_excel(data, report_type, pd.Timestamp("2026-06-10"), context)),
        read_only=True,
    )
    start, end, previous_start, previous_end = report_period(report_type, pd.Timestamp("2026-06-10"))

    assert f"渠道销售{label}" in report
    assert "管理层摘要" in report and "行动跟进表" in report and "复盘要求" in report
    assert "同比数据不足" in report
    assert not tables["channel"].empty and not tables["forecast"].empty
    assert tables["sku"]["stock_qty"].sum() == 60
    assert tables["inventory"]["stock_qty"].sum() == 60
    assert start <= end and previous_start <= previous_end
    assert {"领导摘要", "报告说明", "经营结论", "渠道对比", "SKU贡献", "每日趋势", "库存风险", "销量预测", "异常行动", "行动跟进"}.issubset(
        set(workbook.sheetnames)
    )


def test_period_report_marks_missing_comparison_baseline():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-10"]),
            "channel": ["盒马"],
            "sku_id": ["H1"],
            "sku_name": ["盒马商品"],
            "city": ["上海"],
            "sales_qty": [10],
            "sales_amount": [100],
            "stock_qty": [30],
        }
    )

    report = generate_period_report(data, "weekly", pd.Timestamp("2026-06-10"))

    assert "上一同期无可比销量数据" in report
    assert "上一周无可比数据" in report
    assert "同比数据不足" in report


def test_period_report_pipeline_section_is_opt_in():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-10"]),
            "channel": ["盒马"],
            "sku_id": ["H1"],
            "sku_name": ["盒马商品"],
            "city": ["上海"],
            "sales_qty": [10],
            "sales_amount": [100],
            "stock_qty": [30],
        }
    )
    pipeline_context = {
        "pipeline_records": [
            {
                "Account / Company": "盒马鲜生",
                "Contact Person": "采购经理",
                "Lead Source": "渠道复购",
                "Stage": "谈判",
                "Deal Size": 128000,
                "Probability": 0.78,
                "Expected Close Date": "2026-06-28",
                "Pain Point": "补货不稳定",
                "Decision Maker": "采购负责人",
                "Next Step": "确认下周采购单",
                "Risk": "价格审批延迟",
                "Status Update": "待确认排期",
            }
        ]
    }

    default_report = generate_period_report(data, "daily", pd.Timestamp("2026-06-10"), pipeline_context)
    enabled_report = generate_period_report(
        data,
        "daily",
        pd.Timestamp("2026-06-10"),
        {**pipeline_context, "include_pipeline": True},
    )

    assert "销售 Pipeline 分析" not in default_report
    assert "销售 Pipeline 分析" in enabled_report
    assert "盒马鲜生" in enabled_report


def test_period_report_calculates_year_over_year_growth():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2025-06-10", "2026-06-10"]),
            "channel": ["盒马", "盒马"],
            "sku_id": ["H1", "H1"],
            "sku_name": ["盒马商品", "盒马商品"],
            "city": ["上海", "上海"],
            "sales_qty": [100, 120],
            "sales_amount": [1000, 1200],
            "stock_qty": [30, 40],
        }
    )

    report = generate_period_report(data, "daily", pd.Timestamp("2026-06-10"))
    tables = period_report_tables(data, "daily", pd.Timestamp("2026-06-10"))

    assert "同比去年同期 增长 20.0%" in report
    assert tables["channel"].iloc[0]["yoy_sales_qty"] == 100
    assert tables["channel"].iloc[0]["yoy_qty_change"] == pytest.approx(0.2)


def test_weekly_and_monthly_periods_are_complete_natural_periods():
    weekly = report_period("weekly", pd.Timestamp("2026-06-10"))
    monthly = report_period("monthly", pd.Timestamp("2026-06-10"))

    assert weekly == (
        pd.Timestamp("2026-06-08"),
        pd.Timestamp("2026-06-14 23:59:59.999999999"),
        pd.Timestamp("2026-06-01"),
        pd.Timestamp("2026-06-07 23:59:59.999999999"),
    )
    assert monthly == (
        pd.Timestamp("2026-06-01"),
        pd.Timestamp("2026-06-30 23:59:59.999999999"),
        pd.Timestamp("2026-05-01"),
        pd.Timestamp("2026-05-31 23:59:59.999999999"),
    )


def test_chart_uses_available_inventory_and_bar_for_single_date():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-10", "2026-06-10"]),
            "channel": ["盒马", "小象"],
            "sales_qty": [0, 0],
            "sales_amount": [0, 0],
            "stock_qty": [30, 40],
        }
    )

    metric, note = choose_available_metric(data, "sales_qty")
    chart = trend_chart(data, "channel", metric)
    spec = chart.to_dict()

    assert metric == "stock_qty"
    assert "自动改为展示库存" in note
    assert spec["mark"]["type"] == "bar"


def test_short_period_trend_chart_formats_date_axis_instead_of_timestamp_number():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-08", "2026-06-09", "2026-06-10"]),
            "channel": ["盒马", "盒马", "盒马"],
            "sales_qty": [10, 12, 14],
        }
    )

    spec = trend_chart(data, "channel", "sales_qty").to_dict()

    assert spec["encoding"]["x"]["timeUnit"] == "yearmonthdate"
    assert spec["encoding"]["x"]["type"] == "ordinal"
    assert spec["encoding"]["x"]["axis"]["format"] == "%m-%d"
    assert spec["encoding"]["x"]["scale"]["paddingInner"] == pytest.approx(0.34)
    assert spec["encoding"]["xOffset"]["scale"]["paddingInner"] == pytest.approx(0.16)


def test_sku_trend_uses_vertical_bottom_legend_for_long_names():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-08", "2026-06-08"]),
            "sku_name": [
                "【小象限定】奥兰天天真柠檬沙瓦气泡酒320ml",
                "奥兰 奥太狼干红葡萄酒四支组合装 375ml*4 西班牙进口",
            ],
            "sales_qty": [10, 12],
        }
    )

    spec = trend_chart(data, "sku_name", "sales_qty").to_dict()
    legend = spec["encoding"]["color"]["legend"]

    assert legend["orient"] == "bottom"
    assert legend["columns"] == 1
    assert legend["labelLimit"] >= 520


def test_horizontal_bar_keeps_blank_categories_and_uses_visible_bar_color():
    data = pd.DataFrame(
        {
            "city": ["", None, "上海"],
            "sales_qty": [0, 0, 12],
        }
    )

    chart = horizontal_bar_chart(data, "city", "sales_qty")
    spec = chart.to_dict()
    dataset = next(iter(spec["datasets"].values()))

    assert {row["city"] for row in dataset} == {"未识别城市/区域", "上海"}
    assert spec["mark"]["type"] == "bar"
    assert spec["mark"]["color"] == "#6558C7"


def test_mapping_does_not_treat_inventory_measure_as_store():
    raw = pd.DataFrame(
        columns=["日期", "商品ID", "城市", "商品销售量", "大仓库存数量", "前置站点库存数量"]
    )

    mapping = infer_mapping(raw.columns)

    assert mapping["store"] is None
    assert mapping["stock_qty"] == "前置站点库存数量"


def test_xiaoxiang_inventory_combines_warehouse_station_and_transit_components():
    raw = pd.DataFrame(
        {
            "日期": ["2026-06-10"],
            "商品ID": ["X1"],
            "城市": ["上海"],
            "商品销售量": [5],
            "大仓库存数量": [100],
            "前置站点库存数量": [30],
            "供应商到大仓在途数量": [20],
            "大仓到门店在途数量": [10],
        }
    )

    result = normalize_sales_data(raw, default_channel="小象", source_file="小象.xlsx")

    assert result.iloc[0]["stock_qty"] == 160
    assert "大仓、前置站点及在途库存合计" in result.attrs["warnings"][0]


def test_weather_fetch_and_note_include_high_temperature_and_strong_wind():
    payload = {
        "current": {
            "time": "2026-06-13T12:00",
            "temperature_2m": 34.5,
            "apparent_temperature": 38.0,
            "weather_code": 1,
            "wind_speed_10m": 18,
            "wind_gusts_10m": 25,
        },
        "daily": {
            "temperature_2m_max": [36, 37],
            "temperature_2m_min": [25, 26],
            "weather_code": [1, 2],
            "wind_gusts_10m_max": [40, 70],
        },
    }

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, *_):
            return False

        def read(self):
            import json

            return json.dumps(payload).encode("utf-8")

    rows, errors = fetch_regional_weather(lambda *_args, **_kwargs: FakeResponse())
    note = format_weather_note(rows)

    assert not errors and len(rows) == 2
    assert "上海" in note and "苏州" in note
    assert "高温风险" in note and "强风/台风风险提示" in note


def test_leadership_excel_has_spacious_summary_and_chinese_headers():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-09", "2026-06-10"]),
            "channel": ["盒马", "小象"],
            "sku_id": ["H1", "X1"],
            "sku_name": ["盒马商品", "小象商品"],
            "city": ["上海", "苏州"],
            "sales_qty": [10, 20],
            "sales_amount": [0, 200],
            "stock_qty": [30, 40],
        }
    )

    workbook = load_workbook(BytesIO(export_period_report_excel(data, "daily", pd.Timestamp("2026-06-10"))))
    summary = workbook["领导摘要"]
    channel = workbook["渠道对比"]

    assert summary["A1"].value == "渠道销售日报"
    assert summary.merged_cells.ranges
    assert summary.column_dimensions["E"].width >= 40
    assert summary.row_dimensions[17].height >= 30
    assert [cell.value for cell in channel[1]][:3] == ["渠道", "销量", "销售额"]


def test_period_report_docx_is_clean_business_brief():
    from docx import Document

    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-09", "2026-06-10", "2026-06-10", "2025-06-10"]),
            "channel": ["盒马", "小象", "盒马", "小象"],
            "sku_id": ["H1", "X1", "H2", "X1"],
            "sku_name": ["盒马商品", "小象商品", "盒马新品", "小象商品"],
            "city": ["上海", "苏州", "上海", "苏州"],
            "sales_qty": [10, 20, 15, 10],
            "sales_amount": [120, 200, 0, 100],
            "stock_qty": [30, 40, 25, 50],
        }
    )

    docx_bytes = export_period_report_docx(
        data,
        "daily",
        pd.Timestamp("2026-06-10"),
        {"weather_note": "高温", "holiday_note": "618预热", "activity_note": "渠道满减"},
    )
    document = Document(BytesIO(docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert docx_bytes[:2] == b"PK"
    assert "渠道销售日报" in text
    assert "管理层摘要" in text
    assert "渠道表现" in text
    assert "TOP SKU" in text
    assert "口径说明" in text
    assert "库存使用渠道 + SKU 的最新快照口径" in text
    assert len(document.tables) >= 4


def test_purchase_order_is_identified_and_kept_out_of_sales():
    raw = pd.DataFrame(
        {
            "采购单号": ["PO1", "PO2"],
            "约定到货日期": ["2026-06-10", "2026-06-11"],
            "skuid": ["A1", "A1"],
            "商品名称": ["无糖酸奶200g", "无糖酸奶200g"],
            "采购数量": [120, 80],
            "含税采购金额(元)": [600, 400],
            "业态": ["小象超市", "小象超市"],
        }
    )

    profile = profile_file(raw, "采购订单.xlsx")
    result = normalize_sales_data(
        raw,
        default_channel=profile["channel"],
        mapping=profile["mapping"],
        source_file="采购订单.xlsx",
        data_type=profile["data_type"],
    )

    assert profile["data_type"] == "采购"
    assert profile["channel"] == "小象"
    assert profile["mapping"]["purchase_qty"] == "采购数量"
    assert result["purchase_qty"].sum() == 200
    assert result["sales_qty"].sum() == 0
    assert result["sales_amount"].sum() == 0


def test_grouped_header_purchase_excel_is_read_as_order_file(tmp_path):
    path = tmp_path / "xiaoxiang-purchase.xlsx"
    rows = [
        ["采购单信息", "采购单信息", "商品信息", "商品信息", "履约信息"],
        ["采购单号", "约定到货日期", "skuid", "商品名称", "采购数量"],
        ["PO1", "2026-06-10", "A1", "无糖酸奶200g", 30],
    ]
    pd.DataFrame(rows).to_excel(path, index=False, header=False)

    raw = read_sales_file(path)
    profile = profile_file(raw, path.name)

    assert list(raw.columns) == ["采购单号", "约定到货日期", "skuid", "商品名称", "采购数量"]
    assert profile["data_type"] == "采购"
    assert profile["mapping"]["purchase_qty"] == "采购数量"


def test_purchase_quantity_flows_to_kpis_tables_and_word_report():
    data = pd.DataFrame(
        {
            "date": pd.to_datetime(["2026-06-10", "2026-06-10"]),
            "channel": ["小象", "小象"],
            "source_file": ["sales.xlsx", "purchase.xlsx"],
            "data_type": ["销售", "采购"],
            "sku_id": ["A1", "A1"],
            "sku_name": ["无糖酸奶200g", "无糖酸奶200g"],
            "city": ["上海", "上海"],
            "store": ["上海门店", "上海门店"],
            "sales_qty": [100, 0],
            "sales_amount": [900, 0],
            "purchase_qty": [0, 140],
            "stock_qty": [30, 0],
        }
    )

    kpis = calculate_kpis(data)
    tables = period_report_tables(data, "daily", pd.Timestamp("2026-06-10"))
    report_text = generate_period_report(data, "daily", pd.Timestamp("2026-06-10"))
    docx_bytes = export_period_report_docx(data, "daily", pd.Timestamp("2026-06-10"))

    assert kpis["purchase_qty"] == 140
    assert tables["channel"].iloc[0]["purchase_qty"] == 140
    assert tables["sku"].iloc[0]["purchase_sales_gap"] == 40
    assert "采销匹配分析" in report_text
    assert "采购 140" in report_text

    from docx import Document

    document = Document(BytesIO(docx_bytes))
    doc_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "采销匹配" in doc_text
    assert "本期采购量" in table_text


def test_purchase_orders_only_compare_skus_present_in_sales_data():
    sales = pd.DataFrame(
        {
            "日期": ["2026-06-10"],
            "商品编码": ["A1"],
            "商品名称": ["无糖酸奶200g"],
            "销量": [100],
            "销售额": [900],
        }
    )
    purchase = pd.DataFrame(
        {
            "采购单号": ["PO1", "PO2"],
            "约定到货日期": ["2026-06-10", "2026-06-10"],
            "skuid": ["A1", "B1"],
            "商品名称": ["无糖酸奶200g", "燕麦吐司400g"],
            "采购数量": [140, 999],
        }
    )
    sales_profile = profile_file(sales, "销售.xlsx")
    purchase_profile = profile_file(purchase, "采购订单.xlsx")

    result = combine_sales_files(
        [
            {
                "frame": sales,
                "channel": "小象",
                "data_type": sales_profile["data_type"],
                "mapping": sales_profile["mapping"],
                "source_file": "销售.xlsx",
            },
            {
                "frame": purchase,
                "channel": "小象",
                "data_type": purchase_profile["data_type"],
                "mapping": purchase_profile["mapping"],
                "source_file": "采购订单.xlsx",
            },
        ]
    )

    assert result["purchase_qty"].sum() == 140
    assert set(result.loc[result["data_type"].eq("采购"), "sku_id"]) == {"A1"}
    assert any("未出现在已上传销售数据" in warning for warning in result.attrs["warnings"])


def test_hema_purchase_order_uses_meaningful_sku_id_and_merges_with_sales_report():
    sales = pd.DataFrame(
        {
            "业务日期": [20260610],
            "商品编码": [441588014],
            "商品名称": ["奥兰 奥太狼干红葡萄酒四支组合装 375ml*4 西班牙进口"],
            "经营店名称": ["上海盒马鲜生门店"],
            "销售数量": [745],
            "库存总数": [2940],
        }
    )
    purchase = pd.DataFrame(
        {
            "采购单号": ["HPO1"],
            "要求到货日期": ["2026-06-10 10:00:00"],
            "物理仓": ["成都常温FDC"],
            "货品ID": [441588014],
            "商品名称": ["奥兰 奥太狼干红葡萄酒四支组合装 375ml*4 西班牙进口"],
            "SKU ID": ["--"],
            "采购数量": [126],
            "28天历史销量": [999],
            "含税采购金额(元)": [10080],
            "在仓库存": [18],
        }
    )
    sales_profile = profile_file(sales, "盒马销售.xlsx")
    purchase_profile = profile_file(purchase, "导出采购订单列表.xlsx")

    result = combine_sales_files(
        [
            {
                "frame": sales,
                "channel": sales_profile["channel"],
                "data_type": sales_profile["data_type"],
                "mapping": sales_profile["mapping"],
                "source_file": "盒马销售.xlsx",
            },
            {
                "frame": purchase,
                "channel": purchase_profile["channel"],
                "data_type": purchase_profile["data_type"],
                "mapping": purchase_profile["mapping"],
                "source_file": "导出采购订单列表.xlsx",
            },
        ]
    )
    tables = period_report_tables(result, "daily", pd.Timestamp("2026-06-10"))

    assert purchase_profile["data_type"] == "采购"
    assert purchase_profile["channel"] == "盒马"
    assert purchase_profile["mapping"]["sku_id"] == "货品ID"
    assert purchase_profile["mapping"]["sales_qty"] is None
    assert purchase_profile["mapping"]["sales_amount"] is None
    assert purchase_profile["mapping"]["stock_qty"] is None
    assert result.loc[result["data_type"].eq("采购"), "sku_id"].iloc[0] == "441588014"
    assert result.loc[result["data_type"].eq("采购"), "stock_qty"].sum() == 0
    assert tables["sku"].iloc[0]["sales_qty"] == 745
    assert tables["sku"].iloc[0]["purchase_qty"] == 126
